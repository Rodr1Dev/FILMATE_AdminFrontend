from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs-manual-admin"
CAPTURES = OUT / "capturas"
DOC_CODE = "MA-FILMATE-ADM-001"
VERSION = "1.0"
BASELINE_DATE = "20 de junio de 2026"

FLD_CHAR = "w:fldChar"
FLD_CHAR_TYPE = "w:fldCharType"
TEAM_FILMATE = "Equipo FILMATE"
COURSE_NAME = "Gestión de la Configuración de Software"
BASELINE_LABEL = "Línea base"
VERSION_LABEL = "Versión"
TABLE_DATE = "20/06/2026"

LOGO_CANDIDATES = [
    ROOT / "public" / "logoGrandeFilmate.png",
    ROOT.parent / "FILMATE_UserFrontend" / "public" / "logoGrandeFilmate.png",
    ROOT.parent / "FILMATE_UserFrontend" / "public" / "logoFilmate.png",
    ROOT.parent / "FILMATE_UserFrontend" / "public" / "logo.png",
]


@dataclass
class Figure:
    filename: str
    caption: str
    details: list[str]


class Manual:
    def __init__(self) -> None:
        self.md: list[str] = []
        self.doc = Document()
        self.pdf_story: list = []
        self.figure_number = 0
        self._configure_docx()
        self._configure_pdf()

    def _configure_docx(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.2)

        styles = self.doc.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10.5)
        styles["Normal"].paragraph_format.space_after = Pt(6)
        styles["Normal"].paragraph_format.line_spacing = 1.15

        H1 = "Heading 1"
        H2 = "Heading 2"
        H3 = "Heading 3"
        heading_colors = {
            "Title": RGBColor(15, 23, 42),
            H1: RGBColor(30, 64, 175),
            H2: RGBColor(190, 24, 93),
            H3: RGBColor(15, 118, 110),
        }
        for name, color in heading_colors.items():
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.color.rgb = color

        styles[H1].font.size = Pt(18)
        styles[H1].paragraph_format.space_before = Pt(14)
        styles[H1].paragraph_format.space_after = Pt(8)
        styles[H2].font.size = Pt(14)
        styles[H3].font.size = Pt(11.5)

        for section in self.doc.sections:
            header = section.header.paragraphs[0]
            header.text = f"FILMATE | Manual de administrador | {DOC_CODE} | v{VERSION}"
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.runs[0].font.size = Pt(8)
            header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

            footer = section.footer.paragraphs[0]
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.add_run("Página ")
            run.font.size = Pt(8)
            fld_char1 = OxmlElement(FLD_CHAR)
            fld_char1.set(qn(FLD_CHAR_TYPE), "begin")
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            instr_text.text = "PAGE"
            fld_char2 = OxmlElement(FLD_CHAR)
            fld_char2.set(qn(FLD_CHAR_TYPE), "end")
            run._r.append(fld_char1)
            run._r.append(instr_text)
            run._r.append(fld_char2)
            run.add_text(" de ")
            fld_char3 = OxmlElement(FLD_CHAR)
            fld_char3.set(qn(FLD_CHAR_TYPE), "begin")
            instr_text2 = OxmlElement("w:instrText")
            instr_text2.set(qn("xml:space"), "preserve")
            instr_text2.text = "NUMPAGES"
            fld_char4 = OxmlElement(FLD_CHAR)
            fld_char4.set(qn(FLD_CHAR_TYPE), "end")
            run._r.append(fld_char3)
            run._r.append(instr_text2)
            run._r.append(fld_char4)

    def _configure_pdf(self) -> None:
        styles = getSampleStyleSheet()
        self.pdf_styles = {
            "title": ParagraphStyle(
                "ManualTitle",
                parent=styles["Title"],
                fontName="Helvetica-Bold",
                fontSize=23,
                leading=28,
                textColor=colors.HexColor("#0f172a"),
                alignment=TA_CENTER,
                spaceAfter=16,
            ),
            "subtitle": ParagraphStyle(
                "ManualSubtitle",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=12,
                leading=16,
                textColor=colors.HexColor("#475569"),
                alignment=TA_CENTER,
                spaceAfter=8,
            ),
            "h1": ParagraphStyle(
                "ManualH1",
                parent=styles["Heading1"],
                fontName="Helvetica-Bold",
                fontSize=17,
                leading=21,
                textColor=colors.HexColor("#1e40af"),
                spaceBefore=10,
                spaceAfter=8,
            ),
            "h2": ParagraphStyle(
                "ManualH2",
                parent=styles["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=13,
                leading=17,
                textColor=colors.HexColor("#be185d"),
                spaceBefore=9,
                spaceAfter=6,
            ),
            "h3": ParagraphStyle(
                "ManualH3",
                parent=styles["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=11,
                leading=14,
                textColor=colors.HexColor("#0f766e"),
                spaceBefore=7,
                spaceAfter=4,
            ),
            "body": ParagraphStyle(
                "ManualBody",
                parent=styles["BodyText"],
                fontName="Helvetica",
                fontSize=9.5,
                leading=13.2,
                alignment=TA_JUSTIFY,
                textColor=colors.HexColor("#1e293b"),
                spaceAfter=5,
            ),
            "caption": ParagraphStyle(
                "ManualCaption",
                parent=styles["BodyText"],
                fontName="Helvetica-Oblique",
                fontSize=8.5,
                leading=11,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#475569"),
                spaceAfter=6,
            ),
            "small": ParagraphStyle(
                "ManualSmall",
                parent=styles["BodyText"],
                fontName="Helvetica",
                fontSize=8,
                leading=10,
                textColor=colors.HexColor("#475569"),
            ),
            "callout": ParagraphStyle(
                "ManualCallout",
                parent=styles["BodyText"],
                fontName="Helvetica-Bold",
                fontSize=9.2,
                leading=12.5,
                borderColor=colors.HexColor("#93c5fd"),
                borderWidth=0.8,
                borderPadding=7,
                backColor=colors.HexColor("#eff6ff"),
                textColor=colors.HexColor("#1e3a8a"),
                spaceBefore=5,
                spaceAfter=7,
            ),
        }

    def _get_logo(self) -> Path | None:
        for p in LOGO_CANDIDATES:
            if p.exists():
                return p
        return None

    def cover(self) -> None:
        self.md += [
            "# MANUAL DE ADMINISTRADOR \u2014 FILMATE",
            "",
            "**Panel de gesti\u00f3n administrativa para cartelera, cines, salas, programaci\u00f3n y ventas**",
            "",
            "| Campo | Informaci\u00f3n |",
            "|---|---|",
            f"| C\u00f3digo del documento | {DOC_CODE} |",
            f"| Versi\u00f3n | {VERSION} |",
            f"| L\u00ednea base documental | {BASELINE_DATE} |",
            "| Curso | Gesti\u00f3n de la Configuraci\u00f3n de Software |",
            "| Tipo de entregable | Manual de administrador con evidencia visual |",
            "| Estado | Versi\u00f3n para evaluaci\u00f3n acad\u00e9mica |",
            "| Elaborado por | Equipo FILMATE |",
            "| Docente | REYES HUAMAN, ANITA MARLENE |",
            "",
            "> Este documento describe el frontend administrativo de FILMATE y sus flujos integrados.",
            "",
        ]

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("MANUAL DE ADMINISTRADOR").bold = True
        title.runs[0].font.size = Pt(27)
        title.runs[0].font.color.rgb = RGBColor(15, 23, 42)

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("FILMATE")
        run.bold = True
        run.font.size = Pt(35)
        run.font.color.rgb = RGBColor(239, 68, 68)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Panel de gesti\u00f3n administrativa para cartelera, cines, salas, programaci\u00f3n y ventas").italic = True

        logo = self._get_logo()
        if logo:
            pic = self.doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(str(logo), width=Inches(2.3))

        self.doc.add_paragraph("")
        self._doc_table(
            [
                ["Campo", "Informaci\u00f3n"],
                ["C\u00f3digo", DOC_CODE],
                [VERSION_LABEL, VERSION],
                [BASELINE_LABEL, BASELINE_DATE],
                ["Curso", COURSE_NAME],
                ["Estado", "Versi\u00f3n para evaluaci\u00f3n acad\u00e9mica"],
                ["Elaborado por", TEAM_FILMATE],
                ["Docente", "REYES HUAMAN, ANITA MARLENE"],
            ],
            header=True,
        )
        self.doc.add_page_break()

        self.pdf_story += [
            Spacer(1, 1.5 * cm),
            Paragraph("MANUAL DE ADMINISTRADOR", self.pdf_styles["title"]),
            Paragraph("FILMATE", ParagraphStyle(
                "FilmateTitle",
                parent=self.pdf_styles["title"],
                fontSize=34,
                leading=38,
                textColor=colors.HexColor("#ef4444"),
            )),
            Paragraph(
                "Panel de gesti\u00f3n administrativa para cartelera, cines, salas, programaci\u00f3n y ventas",
                self.pdf_styles["subtitle"],
            ),
            Spacer(1, 1.2 * cm),
        ]
        if logo:
            img = self._pdf_image(logo, max_w=5.3 * cm, max_h=5.3 * cm)
            self.pdf_story += [img, Spacer(1, 0.8 * cm)]
        self.pdf_story.append(
            self._pdf_table(
                [
                    ["Campo", "Informaci\u00f3n"],
                    ["C\u00f3digo", DOC_CODE],
                    [VERSION_LABEL, VERSION],
                    [BASELINE_LABEL, BASELINE_DATE],
                    ["Curso", COURSE_NAME],
                    ["Estado", "Versi\u00f3n para evaluaci\u00f3n acad\u00e9mica"],
                    ["Elaborado por", TEAM_FILMATE],
                    ["Docente", "REYES HUAMAN, ANITA MARLENE"],
                ],
                widths=[4.7 * cm, 10.3 * cm],
            )
        )
        self.pdf_story += [
            Spacer(1, 0.6 * cm),
            PageBreak(),
        ]

    def heading(self, text: str, level: int = 1) -> None:
        self.md += [f"{'#' * level} {text}", ""]
        self.doc.add_heading(text, level=level)
        self.pdf_story.append(Paragraph(text, self.pdf_styles[f"h{level}"]))

    def paragraph(self, text: str, callout: bool = False) -> None:
        prefix = "> " if callout else ""
        self.md += [f"{prefix}{text}", ""]
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        if callout:
            run.bold = True
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EFF6FF")
            p._p.get_or_add_pPr().append(shading)
        self.pdf_story.append(Paragraph(self._escape(text), self.pdf_styles["callout" if callout else "body"]))

    def bullets(self, items: Iterable[str], ordered: bool = False) -> None:
        items = list(items)
        for index, item in enumerate(items, start=1):
            marker = f"{index}." if ordered else "-"
            self.md.append(f"{marker} {item}")
            style = "List Number" if ordered else "List Bullet"
            p = self.doc.add_paragraph(style=style)
            p.add_run(item)
        self.md.append("")

        bullet_type = "1" if ordered else "bullet"
        self.pdf_story.append(
            ListFlowable(
                [
                    ListItem(
                        Paragraph(self._escape(item), self.pdf_styles["body"]),
                        leftIndent=12,
                    )
                    for item in items
                ],
                bulletType=bullet_type,
                leftIndent=22,
                bulletFontName="Helvetica",
                bulletFontSize=8,
                spaceAfter=6,
            )
        )

    def table(self, rows: list[list[str]]) -> None:
        self.md.append("| " + " | ".join(rows[0]) + " |")
        self.md.append("|" + "|".join(["---"] * len(rows[0])) + "|")
        for row in rows[1:]:
            self.md.append("| " + " | ".join(row) + " |")
        self.md.append("")
        self._doc_table(rows, header=True)
        col_count = len(rows[0])
        usable = 15.2 * cm
        widths = [usable / col_count] * col_count
        self.pdf_story += [self._pdf_table(rows, widths=widths), Spacer(1, 0.2 * cm)]

    def figure(self, figure: Figure) -> None:
        self.figure_number += 1
        path = CAPTURES / figure.filename
        caption = f"Figura {self.figure_number}. {figure.caption}"
        rel = Path("capturas") / figure.filename
        self.md += [f"![{caption}]({rel.as_posix()})", "", f"*{caption}*", ""]
        if figure.details:
            self.md.append("Elementos y lectura de la pantalla:")
            self.md.append("")
            for item in figure.details:
                self.md.append(f"- {item}")
            self.md.append("")

        if path.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width, height = self._doc_image_size(path)
            p.add_run().add_picture(str(path), width=width, height=height)
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(caption)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(71, 85, 105)
            for item in figure.details:
                bp = self.doc.add_paragraph(style="List Bullet")
                bp.add_run(item)

            pdf_img = self._pdf_image(path, max_w=17.1 * cm, max_h=17.2 * cm)
            block = [pdf_img, Spacer(1, 0.12 * cm), Paragraph(caption, self.pdf_styles["caption"])]
            self.pdf_story += [KeepTogether(block)]
            if figure.details:
                self.pdf_story.append(
                    ListFlowable(
                        [
                            ListItem(Paragraph(self._escape(item), self.pdf_styles["small"]), leftIndent=10)
                            for item in figure.details
                        ],
                        bulletType="bullet",
                        leftIndent=20,
                        bulletFontSize=7,
                        spaceAfter=6,
                    )
                )

    def page_break(self) -> None:
        self.md += ['<div style="page-break-after: always;"></div>', ""]
        self.doc.add_page_break()
        self.pdf_story.append(PageBreak())

    def save(self) -> None:
        OUT.mkdir(parents=True, exist_ok=True)
        md_path = OUT / "Manual_de_Administrador_FILMATE.md"
        docx_path = OUT / "Manual_de_Administrador_FILMATE.docx"
        pdf_path = OUT / "Manual_de_Administrador_FILMATE.pdf"

        md_path.write_text("\n".join(self.md), encoding="utf-8")

        self.doc.save(str(docx_path))

        pdf = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            rightMargin=1.7 * cm,
            leftMargin=1.7 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.7 * cm,
            title="Manual de Administrador FILMATE",
            author="Equipo FILMATE",
            subject="Gesti\u00f3n de la Configuraci\u00f3n de Software",
        )
        pdf.build(self.pdf_story, onFirstPage=self._pdf_page, onLaterPages=self._pdf_page)

        print(md_path)
        print(docx_path)
        print(pdf_path)

    def _doc_table(self, rows: list[list[str]], header: bool = True) -> None:
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for r_index, row in enumerate(rows):
            for c_index, value in enumerate(row):
                cell = table.cell(r_index, c_index)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = str(value)
                self._style_doc_cell(cell, header, r_index)
        self.doc.add_paragraph("")

    def _style_doc_cell(self, cell, header: bool, r_index: int) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if header and r_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if header and r_index == 0:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "1E3A8A")
            cell._tc.get_or_add_tcPr().append(shading)

    def _pdf_table(self, rows: list[list[str]], widths: list[float]) -> Table:
        data = []
        for row_index, row in enumerate(rows):
            style = ParagraphStyle(
                f"TableCell{row_index}",
                parent=self.pdf_styles["small"],
                textColor=colors.white if row_index == 0 else colors.HexColor("#1e293b"),
                fontName="Helvetica-Bold" if row_index == 0 else "Helvetica",
                leading=10,
            )
            data.append([Paragraph(self._escape(str(cell)), style) for cell in row])
        table = Table(data, colWidths=widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]
            )
        )
        return table

    def _doc_image_size(self, path: Path) -> tuple:
        with PILImage.open(path) as image:
            px_w, px_h = image.size
        max_w = 6.45
        max_h = 7.1
        scale = min(max_w / px_w, max_h / px_h)
        return Inches(px_w * scale), Inches(px_h * scale)

    def _pdf_image(self, path: Path, max_w: float, max_h: float) -> Image:
        with PILImage.open(path) as image:
            px_w, px_h = image.size
        scale = min(max_w / px_w, max_h / px_h)
        return Image(str(path), width=px_w * scale, height=px_h * scale)

    def _pdf_page(self, canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(1.7 * cm, A4[1] - 1.05 * cm, f"FILMATE | {DOC_CODE} | v{VERSION}")
        canvas.drawRightString(A4[0] - 1.7 * cm, 0.9 * cm, f"P\u00e1gina {doc.page}")
        canvas.restoreState()

    @staticmethod
    def _escape(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )


def add_control_and_scope(manual: Manual) -> None:
    manual.heading("Control del documento", 1)
    manual.table(
        [
            [VERSION_LABEL, "Fecha", "Descripci\u00f3n del cambio", "Responsable"],
            ["0.1", TABLE_DATE, "Inventario funcional y planificaci\u00f3n de capturas.", TEAM_FILMATE],
            ["0.9", TABLE_DATE, "Ejecuci\u00f3n de 26 evidencias visuales y validaci\u00f3n de flujos.", TEAM_FILMATE],
            ["1.0", TABLE_DATE, "Emisi\u00f3n del manual acad\u00e9mico en Markdown, Word y PDF.", TEAM_FILMATE],
        ]
    )
    manual.paragraph(
        f"Criterio de configuraci\u00f3n documental: este manual se identifica mediante el c\u00f3digo "
        f"{DOC_CODE}; la versi\u00f3n {VERSION} queda asociada a la l\u00ednea base visual del {BASELINE_DATE}. "
        "Cualquier cambio en rutas, etiquetas, reglas de validaci\u00f3n, flujos de administraci\u00f3n o dise\u00f1o de pantallas "
        "debe originar una nueva revisi\u00f3n del documento y la sustituci\u00f3n de las capturas afectadas.",
        callout=True,
    )

    manual.heading("Contenido", 1)
    manual.bullets(
        [
            "1. Introducci\u00f3n, objetivo y alcance.",
            "2. Descripci\u00f3n del sistema, perfiles y requisitos de uso.",
            "3. Convenciones visuales y mapa de navegaci\u00f3n.",
            "4. Acceso e inicio de sesi\u00f3n.",
            "5. Barra lateral y navegaci\u00f3n.",
            "6. Dashboard principal: estad\u00edsticas y gr\u00e1ficos.",
            "7. Dashboard principal: \u00faltimas transacciones.",
            "8. Cat\u00e1logo de pel\u00edculas.",
            "9. Cines y salas (CRUD).",
            "10. Programaci\u00f3n y funciones.",
            "11. Ventas y tickets.",
            "12. Reembolsos y devoluciones.",
            "13. Validaci\u00f3n de entrada.",
            "14. Ayuda y soporte.",
            "15. Cierre de sesi\u00f3n.",
            "16. Soluci\u00f3n de problemas.",
            "17. Seguridad, privacidad y buenas pr\u00e1cticas.",
            "18. Trazabilidad, criterios de aceptaci\u00f3n y gesti\u00f3n del manual.",
            "19. Glosario y gu\u00eda r\u00e1pida.",
        ]
    )

    manual.heading("1. Introducci\u00f3n", 1)
    manual.heading("1.1 Objetivo", 2)
    manual.paragraph(
        "Orientar al administrador en el uso seguro y correcto del panel de gesti\u00f3n FILMATE, desde el acceso inicial "
        "hasta la administraci\u00f3n de pel\u00edculas, cines, salas, programaci\u00f3n, ventas, reembolsos y soporte. "
        "Cada procedimiento incluye precondiciones, acciones, resultado esperado y evidencia visual."
    )
    manual.heading("1.2 Alcance", 2)
    manual.paragraph(
        "El manual cubre el frontend administrativo: inicio de sesi\u00f3n, dashboard con estad\u00edsticas y gr\u00e1ficos, "
        "cat\u00e1logo de pel\u00edculas con filtros, gesti\u00f3n de cines y salas (CRUD), programaci\u00f3n de funciones, "
        "ventas y tickets con filtros y detalle, flujo completo de reembolsos (solicitud, revisi\u00f3n, resoluci\u00f3n), "
        "validaci\u00f3n de entrada y m\u00f3dulo de ayuda y soporte."
    )
    manual.paragraph(
        "No cubre el frontend de usuario (cartelera p\u00fablica, reserva, dulcer\u00eda, perfil social) ni la configuraci\u00f3n "
        "interna del servidor, base de datos o infraestructura; esas operaciones corresponden al equipo de desarrollo."
    )
    manual.heading("1.3 Base de la evidencia", 2)
    manual.paragraph(
        "Las capturas se obtuvieron el 20 de junio de 2026 sobre la aplicaci\u00f3n React/Vite del repositorio "
        "FILMATE_AdminFrontend. Se utiliz\u00f3 un entorno de demostraci\u00f3n local controlado con la misma estructura "
        "de datos y contratos de API. No se utilizaron datos financieros reales."
    )

    manual.heading("2. Descripci\u00f3n general del sistema", 1)
    manual.paragraph(
        "FILMATE Admin es un panel de gesti\u00f3n web orientado al personal administrativo de cines. "
        "Integra la administraci\u00f3n de pel\u00edculas, locales, salas, horarios, ventas y soporte al cliente."
    )
    manual.heading("2.1 Perfiles de acceso", 2)
    manual.table(
        [
            ["Perfil", "Capacidades principales", "Restricciones"],
            ["Administrador", "Gesti\u00f3n completa: cat\u00e1logos, programaci\u00f3n, ventas, reembolsos y soporte.", "Debe iniciar sesi\u00f3n con credenciales administrativas."],
            ["Validador", "Validar entradas en punto de atenci\u00f3n.", "No accede a cat\u00e1logos, ventas ni reembolsos."],
        ]
    )
    manual.heading("2.2 Requisitos m\u00ednimos", 2)
    manual.bullets(
        [
            "Equipo de escritorio, port\u00e1til o tableta con navegador moderno.",
            "Google Chrome, Microsoft Edge o Mozilla Firefox actualizado.",
            "Resoluci\u00f3n recomendada de escritorio: 1366 \u00d7 768 o superior.",
            "Conexi\u00f3n de red estable con acceso al frontend y al servicio API.",
            "Credenciales de administrador proporcionadas por el \u00e1rea de TI.",
            "JavaScript y almacenamiento local habilitados en el navegador.",
        ]
    )

    manual.heading("3. Convenciones y navegaci\u00f3n", 1)
    manual.heading("3.1 Convenciones visuales", 2)
    manual.table(
        [
            ["Elemento", "Significado"],
            ["Bot\u00f3n rojo", "Acci\u00f3n destructiva (eliminar, cancelar) o cierre de sesi\u00f3n."],
            ["Bot\u00f3n azul", "Acci\u00f3n de navegaci\u00f3n, edici\u00f3n o consulta."],
            ["Bot\u00f3n verde", "Confirmaci\u00f3n o creaci\u00f3n de recurso."],
            ["Mensaje rojo", "Error de validaci\u00f3n o imposibilidad de completar una operaci\u00f3n."],
            ["Mensaje verde", "Operaci\u00f3n completada correctamente."],
            ["Panel lateral", "Navegaci\u00f3n entre m\u00f3dulos del sistema."],
        ]
    )
    manual.heading("3.2 Mapa de navegaci\u00f3n", 2)
    manual.paragraph(
        "La barra lateral izquierda permite navegar entre los m\u00f3dulos: Dashboard, Cat\u00e1logo, "
        "Cines y Salas, Programaci\u00f3n, Ventas y Tickets, Validaci\u00f3n de Entrada y Ayuda/Soporte."
    )
    manual.paragraph(
        "Dashboard \u2192 Cat\u00e1logo \u2192 Cines y Salas \u2192 Programaci\u00f3n \u2192 Ventas \u2192 Validaci\u00f3n \u2192 Ayuda.",
        callout=True,
    )


def add_access_and_sidebar(manual: Manual) -> None:
    manual.heading("4. Acceso al sistema", 1)
    manual.heading("4.1 Iniciar sesi\u00f3n como administrador", 2)
    manual.paragraph("Precondici\u00f3n: el administrador debe disponer de una cuenta activa con rol administrativo.")
    manual.bullets(
        [
            "Abra la URL del panel administrativo en el navegador.",
            "Escriba el correo electr\u00f3nico registrado.",
            "Escriba la contrase\u00f1a; el valor se oculta por seguridad.",
            "Seleccione \u00abIniciar sesi\u00f3n\u00bb o presione Enter.",
            "Espere el mensaje de \u00e9xito y la redirecci\u00f3n autom\u00e1tica al Dashboard.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "01-Inicio-sesion-admin.png",
            "Pantalla de inicio de sesi\u00f3n del panel administrativo.",
            [
                "El campo de correo exige un formato que incluya el car\u00e1cter @.",
                "La contrase\u00f1a se muestra en modo oculto.",
                "Las credenciales deben corresponder a un perfil administrativo.",
            ],
        )
    )

    manual.heading("4.2 Validaciones de acceso", 2)
    manual.paragraph(
        "Si falta un dato, el correo no tiene formato v\u00e1lido o las credenciales no coinciden, "
        "FILMATE presenta un mensaje de error. Corrija el dato indicado y vuelva a intentar."
    )

    manual.heading("5. Barra lateral y navegaci\u00f3n", 1)
    manual.heading("5.1 Estructura de la barra lateral", 2)
    manual.paragraph(
        "Una vez dentro del panel, la barra lateral izquierda agrupa el acceso a todos los m\u00f3dulos "
        "administrativos. El m\u00f3dulo activo se resalta visualmente."
    )
    manual.figure(
        Figure(
            "02-Barra-lateral-navegable.png",
            "Barra lateral con todos los m\u00f3dulos administrativos.",
            [
                "Dashboard: estad\u00edsticas y gr\u00e1ficos del negocio.",
                "Cat\u00e1logo: gesti\u00f3n de pel\u00edculas.",
                "Cines y Salas: administraci\u00f3n de locales y salas.",
                "Programaci\u00f3n: asignaci\u00f3n de funciones.",
                "Ventas y Tickets: consulta y gesti\u00f3n de transacciones.",
                "Validaci\u00f3n de Entrada: escaneo de tickets.",
                "Ayuda y Soporte: gesti\u00f3n de consultas.",
            ],
        )
    )
    manual.heading("5.2 Cerrar sesi\u00f3n desde la barra", 2)
    manual.bullets(
        [
            "Ubique el bot\u00f3n de cerrar sesi\u00f3n al final de la barra lateral.",
            "Selecci\u00f3nelo y confirme la acci\u00f3n en el di\u00e1logo emergente.",
            "Espere la redirecci\u00f3n a la pantalla de inicio de sesi\u00f3n.",
        ],
        ordered=True,
    )


def add_dashboard(manual: Manual) -> None:
    manual.heading("6. Dashboard principal", 1)
    manual.heading("6.1 Estad\u00edsticas y gr\u00e1ficos", 2)
    manual.paragraph(
        "El Dashboard presenta un resumen visual del negocio: tarjetas con m\u00e9tricas clave y gr\u00e1ficos "
        "de ingresos. El administrador puede filtrar por per\u00edodo (hoy, semana, mes, mes anterior)."
    )
    manual.figure(
        Figure(
            "03-Dashboard-Principal-estadisticas-y-graficos.png",
            "Dashboard con tarjetas de m\u00e9tricas y gr\u00e1fico de ingresos por categor\u00eda.",
            [
                "Tarjetas superiores: total de transacciones, ingresos totales, productos m\u00e1s vendidos.",
                "Selector de per\u00edodo (hoy, semana, mes, mes anterior) para filtrar datos.",
                "Gr\u00e1fico circular de ingresos por categor\u00eda.",
                "Gr\u00e1fico de barras para comparativas de rendimiento.",
            ],
        )
    )

    manual.heading("6.2 \u00daltimas transacciones", 2)
    manual.paragraph(
        "Debajo de los gr\u00e1ficos, una tabla muestra las transacciones m\u00e1s recientes registradas en el sistema."
    )
    manual.figure(
        Figure(
            "04-Dashboard-Principal-ultimas-transacciones.png",
            "Tabla de \u00faltimas transacciones en el Dashboard.",
            [
                "Columnas: ID, usuario, pel\u00edcula, fecha, monto, estado.",
                "La tabla se actualiza autom\u00e1ticamente al cambiar el filtro de per\u00edodo.",
                "Seleccione una fila para ver el detalle completo de la transacci\u00f3n.",
            ],
        )
    )


def add_catalog(manual: Manual) -> None:
    manual.heading("7. Cat\u00e1logo de pel\u00edculas", 1)
    manual.heading("7.1 Vista general del cat\u00e1logo", 2)
    manual.paragraph(
        "El m\u00f3dulo de Cat\u00e1logo lista todas las pel\u00edculas registradas en el sistema. "
        "Cada pel\u00edcula muestra p\u00f3ster, t\u00edtulo, g\u00e9nero, clasificaci\u00f3n y estado."
    )
    manual.figure(
        Figure(
            "05-Catalogo-de-peliculas.png",
            "Listado general del cat\u00e1logo de pel\u00edculas.",
            [
                "Cada tarjeta incluye p\u00f3ster, t\u00edtulo, g\u00e9nero, duraci\u00f3n, clasificaci\u00f3n y estado.",
                "Botones de acci\u00f3n: editar y eliminar por pel\u00edcula.",
                "Bot\u00f3n \u00abAgregar Pel\u00edcula\u00bb para incorporar nuevos t\u00edtulos.",
            ],
        )
    )

    manual.heading("7.2 Filtros del cat\u00e1logo", 2)
    manual.bullets(
        [
            "Use el campo de b\u00fasqueda para encontrar pel\u00edculas por t\u00edtulo.",
            "Filtre por g\u00e9nero, clasificaci\u00f3n o estado (activa/inactiva).",
            "Seleccione \u00abLimpiar filtros\u00bb para restablecer la vista completa.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "06-Catalago-de-peliculas-FIltros.png",
            "Cat\u00e1logo filtrado por criterios de b\u00fasqueda.",
            [
                "Los filtros se aplican en tiempo real sobre el listado.",
                "Combine m\u00faltiples criterios para refinar la b\u00fasqueda.",
                "Si no hay coincidencias, ajuste los filtros aplicados.",
            ],
        )
    )


def add_cinemas_and_rooms(manual: Manual) -> None:
    manual.heading("8. Cines y salas", 1)
    manual.heading("8.1 Vista general de cines", 2)
    manual.paragraph(
        "El m\u00f3dulo de Cines y Salas permite administrar los locales y sus respectivas salas de exhibici\u00f3n."
    )
    manual.figure(
        Figure(
            "07-Cines-y-salas.png",
            "Listado de cines registrados en el sistema.",
            [
                "Cada cine muestra nombre, direcci\u00f3n, tel\u00e9fono y n\u00famero de salas.",
                "Botones de acci\u00f3n: editar cine, ver salas, eliminar.",
                "Bot\u00f3n \u00abAgregar Cine\u00bb para registrar un nuevo local.",
            ],
        )
    )

    manual.heading("8.2 Agregar un cine", 2)
    manual.bullets(
        [
            "Seleccione \u00abAgregar Cine\u00bb en la parte superior del listado.",
            "Complete el formulario: nombre, direcci\u00f3n, tel\u00e9fono y estado.",
            "Seleccione \u00abGuardar\u00bb para registrar el nuevo cine.",
            "Confirme que el cine aparece en el listado principal.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "08-Cines-y-salas-Agregar-CIne.png",
            "Formulario para agregar un nuevo cine.",
            [
                "Todos los campos obligatorios est\u00e1n marcados.",
                "El tel\u00e9fono debe tener formato v\u00e1lido.",
                "El estado permite activar o desactivar el cine.",
            ],
        )
    )

    manual.heading("8.3 Agregar una sala", 2)
    manual.bullets(
        [
            "Seleccione un cine del listado y acceda a sus salas.",
            "Seleccione \u00abAgregar Sala\u00bb.",
            "Complete nombre de sala, tipo y capacidad de asientos.",
            "Guarde y verifique que la sala aparezca en el listado del cine.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "09-Cines-y-salas-Agregar-Sala.png",
            "Formulario de registro de nueva sala.",
            [
                "El nombre debe ser \u00fanico dentro del mismo cine.",
                "El tipo de sala puede ser 2D, 3D o VIP.",
                "La capacidad define el n\u00famero de asientos disponibles.",
            ],
        )
    )

    manual.heading("8.4 Editar una sala", 2)
    manual.bullets(
        [
            "Localice la sala en el listado del cine correspondiente.",
            "Seleccione el icono de edici\u00f3n.",
            "Modifique los campos necesarios: nombre, tipo, capacidad.",
            "Guarde los cambios y verifique la actualizaci\u00f3n.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "10-Cines-y-salas-Editar-Sala.png",
            "Formulario de edici\u00f3n de sala con datos precargados.",
            [
                "Los campos se cargan con la informaci\u00f3n actual de la sala.",
                "Es posible cambiar el tipo o capacidad seg\u00fan necesidad.",
                "La actualizaci\u00f3n no afecta funciones futuras ya programadas.",
            ],
        )
    )

    manual.heading("8.5 Eliminar una sala", 2)
    manual.paragraph("Precondici\u00f3n: la sala no debe tener funciones activas programadas.")
    manual.bullets(
        [
            "Localice la sala en el listado del cine.",
            "Seleccione el icono de eliminaci\u00f3n.",
            "Confirme la acci\u00f3n en el di\u00e1logo emergente.",
            "Verifique que la sala desaparezca del listado.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "11-Cines-y-salas-Eliminar-Sala.png",
            "Di\u00e1logo de confirmaci\u00f3n para eliminar una sala.",
            [
                "El sistema solicita confirmaci\u00f3n antes de eliminar.",
                "Si la sala tiene funciones activas, se mostrar\u00e1 una advertencia.",
                "La eliminaci\u00f3n es permanente y no se puede deshacer.",
            ],
        )
    )


def add_scheduling(manual: Manual) -> None:
    manual.heading("9. Programaci\u00f3n", 1)
    manual.heading("9.1 Vista general de programaci\u00f3n", 2)
    manual.paragraph(
        "El m\u00f3dulo de Programaci\u00f3n permite asignar pel\u00edculas a salas con horarios espec\u00edficos."
    )
    manual.figure(
        Figure(
            "12-Programacion.png",
            "Vista general de la programaci\u00f3n de funciones.",
            [
                "Listado de funciones programadas con pel\u00edcula, cine, sala, fecha y hora.",
                "Botones para agregar, editar y eliminar funciones.",
                "Filtros por cine, pel\u00edcula o fecha para localizar funciones.",
            ],
        )
    )

    manual.heading("9.2 Elegir fecha y filtros", 2)
    manual.bullets(
        [
            "Use el selector de fecha para visualizar funciones de un d\u00eda espec\u00edfico.",
            "Filtre por cine o pel\u00edcula para reducir los resultados.",
            "Seleccione una fecha futura para planificar nueva programaci\u00f3n.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "13-Programacion-Elegir-Fecha.png",
            "Selector de fecha para filtrar la programaci\u00f3n.",
            [
                "El calendario permite navegar entre meses.",
                "Las fechas con funciones se muestran con un indicador visual.",
                "Seleccione cualquier fecha para ver su programaci\u00f3n.",
            ],
        )
    )

    manual.heading("9.3 Visualizar y gestionar funciones", 2)
    manual.bullets(
        [
            "Revise las funciones listadas para la fecha seleccionada.",
            "Cada funci\u00f3n muestra pel\u00edcula, sala, hora de inicio y precio.",
            "Use \u00abAgregar Funci\u00f3n\u00bb para crear un nuevo horario.",
            "Edite o elimine funciones seg\u00fan sea necesario.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "14-Programacion-Ver-Funciones.png",
            "Funciones programadas para una fecha espec\u00edfica.",
            [
                "Las funciones se agrupan por cine y sala.",
                "El precio base de la funci\u00f3n se muestra en la lista.",
                "Evite superponer funciones en la misma sala y horario.",
            ],
        )
    )


def add_sales_and_tickets(manual: Manual) -> None:
    manual.heading("10. Ventas y tickets", 1)
    manual.heading("10.1 Vista general de ventas", 2)
    manual.paragraph(
        "El m\u00f3dulo de Ventas y Tickets lista todas las transacciones realizadas a trav\u00e9s de la plataforma."
    )
    manual.figure(
        Figure(
            "15-Ventas-y-Tickets.png",
            "Listado general de ventas y tickets.",
            [
                "Cada fila representa una transacci\u00f3n con ID, usuario, pel\u00edcula y monto.",
                "El estado indica si la venta est\u00e1 activa, reembolsada o en proceso.",
                "Botones de acci\u00f3n: ver detalle y procesar reembolso.",
            ],
        )
    )

    manual.heading("10.2 Filtros de b\u00fasqueda", 2)
    manual.bullets(
        [
            "Filtre por rango de fechas para acotar los resultados.",
            "Busque por ID de transacci\u00f3n o nombre de usuario.",
            "Filtre por estado: activa, reembolsada o en proceso.",
            "Seleccione \u00abLimpiar filtros\u00bb para restablecer la vista completa.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "16-Ventas-y-Tickets-FIltros.png",
            "Ventas filtradas por criterios de b\u00fasqueda.",
            [
                "Los filtros se aplican en el servidor para optimizar la b\u00fasqueda.",
                "Combine fecha, usuario y estado para localizar una transacci\u00f3n espec\u00edfica.",
                "El resultado muestra solo las transacciones que cumplen todos los criterios.",
            ],
        )
    )

    manual.heading("10.3 Detalle de ticket", 2)
    manual.paragraph(
        "Al seleccionar una transacci\u00f3n, se muestra el detalle completo del ticket, incluyendo "
        "pel\u00edcula, cine, sala, fecha, asientos, productos de dulcer\u00eda y m\u00e9todo de pago."
    )
    manual.figure(
        Figure(
            "17-Ventas-y-Tickets-Detalle-de-Ticket.png",
            "Detalle completo de un ticket de venta.",
            [
                "Informaci\u00f3n de la funci\u00f3n: pel\u00edcula, cine, sala, fecha y hora.",
                "Listado de asientos adquiridos con fila y n\u00famero.",
                "Productos de dulcer\u00eda incluidos en la compra.",
                "M\u00e9todo de pago, monto total y estado de la transacci\u00f3n.",
            ],
        )
    )


def add_refunds(manual: Manual) -> None:
    manual.heading("11. Reembolsos y devoluciones", 1)
    manual.heading("11.1 Solicitar reembolso", 2)
    manual.paragraph(
        "El administrador puede iniciar un proceso de reembolso desde el detalle del ticket."
    )
    manual.figure(
        Figure(
            "18-Ventas-y-Tickets-Solicitar-Reembolso.png",
            "Formulario de solicitud de reembolso.",
            [
                "Seleccione el motivo del reembolso de la lista predefinida.",
                "Agregue un comentario opcional para justificar la solicitud.",
                "Confirme la solicitud para iniciar el proceso.",
            ],
        )
    )

    manual.heading("11.2 Devoluciones y reembolsos", 2)
    manual.paragraph(
        "El sistema mantiene un registro de todas las solicitudes de reembolso con su estado actual."
    )
    manual.figure(
        Figure(
            "19-Ventas-y-Tickets-Devoluciones-y-Reembolso.png",
            "Listado de devoluciones y reembolsos.",
            [
                "Cada solicitud muestra ID, usuario, monto, motivo y estado.",
                "Estados: pendiente, aprobado, rechazado.",
                "Seleccione una solicitud para ver su detalle o resolverla.",
            ],
        )
    )

    manual.heading("11.3 Ver detalle de solicitud de reembolso", 2)
    manual.bullets(
        [
            "Localice la solicitud en el listado de reembolsos.",
            "Seleccione la solicitud para ver su detalle completo.",
            "Revise la informaci\u00f3n del ticket original y el motivo indicado.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "20-Ventas-y-Tickets-Ver-Detalle-Reembolso.png",
            "Detalle completo de una solicitud de reembolso.",
            [
                "Muestra el ticket original asociado a la solicitud.",
                "Incluye el motivo y comentario del solicitante.",
                "Botones para aprobar o rechazar la solicitud.",
            ],
        )
    )

    manual.heading("11.4 Resolver una solicitud de reembolso", 2)
    manual.bullets(
        [
            "Acceda al detalle de la solicitud pendiente.",
            "Revise la informaci\u00f3n y el motivo presentado.",
            "Seleccione \u00abAprobar\u00bb o \u00abRechazar\u00bb seg\u00fan corresponda.",
            "Agregue un comentario sobre la resoluci\u00f3n si es necesario.",
            "Confirme la acci\u00f3n para actualizar el estado de la solicitud.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "21-Ventas-y-Tickets-Resolver-Solicitud.png",
            "Resoluci\u00f3n de una solicitud de reembolso.",
            [
                "El administrador puede aprobar o rechazar la solicitud.",
                "Al aprobar, el sistema procesa la devoluci\u00f3n del monto.",
                "El ticket original se marca como reembolsado.",
            ],
        )
    )

    manual.heading("11.5 Solicitud resuelta", 2)
    manual.paragraph(
        "Una vez procesada, la solicitud cambia su estado a aprobado o rechazado y se registra "
        "la fecha de resoluci\u00f3n."
    )
    manual.figure(
        Figure(
            "22-Ventas-y-Tickets-Solicitud-Resuelta.png",
            "Solicitud de reembolso resuelta exitosamente.",
            [
                "El estado refleja la decisi\u00f3n del administrador.",
                "La fecha de resoluci\u00f3n queda registrada en el sistema.",
                "El ticket asociado se actualiza seg\u00fan corresponda.",
            ],
        )
    )

    manual.heading("11.6 Ver detalle de solicitud resuelta", 2)
    manual.bullets(
        [
            "Acceda al listado de reembolsos y localice la solicitud resuelta.",
            "Seleccione la solicitud para ver su detalle.",
            "Revise el estado final, comentarios y fecha de resoluci\u00f3n.",
        ],
        ordered=True,
    )
    manual.figure(
        Figure(
            "23-Ventas-y-Tickets-Ver-detalle-de-de-solcitud-resuelta.png",
            "Detalle de una solicitud de reembolso ya resuelta.",
            [
                "Muestra el estado final: aprobado o rechazado.",
                "Incluye la fecha y comentario de resoluci\u00f3n.",
                "El ticket original refleja el cambio de estado.",
            ],
        )
    )


def add_validation(manual: Manual) -> None:
    manual.heading("12. Validaci\u00f3n de entrada", 1)
    manual.heading("12.1 Escanear ticket", 2)
    manual.paragraph(
        "El m\u00f3dulo de Validaci\u00f3n permite escanear el c\u00f3digo QR del ticket para verificar su autenticidad "
        "y permitir el ingreso a la sala."
    )
    manual.figure(
        Figure(
            "24-Validacion-de-entrada.png",
            "Pantalla de validaci\u00f3n de entrada con escaneo de QR.",
            [
                "Active la c\u00e1mara para escanear el c\u00f3digo QR del ticket.",
                "El sistema muestra la informaci\u00f3n de la funci\u00f3n al detectar un ticket v\u00e1lido.",
                "Indicador visual del estado: v\u00e1lido, usado o inv\u00e1lido.",
            ],
        )
    )

    manual.heading("12.2 Validaci\u00f3n de entrada doble", 2)
    manual.paragraph(
        "Para grupos o compras m\u00faltiples, el sistema permite validar varias entradas "
        "de una misma transacci\u00f3n de forma secuencial."
    )
    manual.figure(
        Figure(
            "25-Validacion-de-entrada-Doble.png",
            "Validaci\u00f3n de m\u00faltiples entradas de una misma transacci\u00f3n.",
            [
                "El sistema muestra las entradas restantes por validar.",
                "Cada validaci\u00f3n registra la hora de ingreso.",
                "Una vez validadas todas, la transacci\u00f3n se marca como completa.",
            ],
        )
    )


def add_support(manual: Manual) -> None:
    manual.heading("13. Ayuda y soporte", 1)
    manual.heading("13.1 Gesti\u00f3n de consultas", 2)
    manual.paragraph(
        "El m\u00f3dulo de Ayuda y Soporte permite al administrador gestionar las consultas "
        "enviadas por los usuarios desde el frontend p\u00fablico."
    )
    manual.figure(
        Figure(
            "26-Ayuda-y-soporte.png",
            "Panel de gesti\u00f3n de consultas de ayuda y soporte.",
            [
                "Listado de consultas recibidas con asunto, usuario y fecha.",
                "Estado de cada consulta: pendiente, en proceso, resuelto.",
                "Seleccione una consulta para ver el detalle y responder.",
                "Historial de respuestas asociadas a cada consulta.",
            ],
        )
    )

    manual.heading("13.2 Responder una consulta", 2)
    manual.bullets(
        [
            "Seleccione la consulta del listado para abrir el detalle.",
            "Revise el mensaje del usuario y la informaci\u00f3n asociada.",
            "Redacte una respuesta en el campo de texto.",
            "Seleccione \u00abEnviar respuesta\u00bb para notificar al usuario.",
            "Actualice el estado de la consulta seg\u00fan corresponda.",
        ],
        ordered=True,
    )


def add_support_and_gcs(manual: Manual) -> None:
    manual.heading("14. Soluci\u00f3n de problemas", 1)
    manual.table(
        [
            ["Problema", "Causa probable", "Soluci\u00f3n del administrador"],
            ["No cargan estad\u00edsticas", "Backend no disponible o red interrumpida.", "Comprobar conexi\u00f3n, esperar y recargar la p\u00e1gina."],
            ["No aparecen pel\u00edculas", "Cat\u00e1logo vac\u00edo o filtros restrictivos.", "Verificar el cat\u00e1logo y limpiar filtros."],
            ["No se puede agregar cine", "Campos obligatorios incompletos.", "Completar todos los campos requeridos."],
            ["Error al eliminar sala", "La sala tiene funciones activas.", "Revisar y eliminar las funciones primero."],
            ["No se puede crear funci\u00f3n", "Superposici\u00f3n de horarios en la misma sala.", "Elegir un horario disponible."],
            ["Solicitud de reembolso no aparece", "Filtros de estado activos.", "Limpiar filtros y verificar el listado completo."],
            ["Validaci\u00f3n de QR falla", "Ticket ya usado, vencido o c\u00f3digo inv\u00e1lido.", "Verificar manualmente el ID de transacci\u00f3n."],
            ["No se env\u00edan respuestas de soporte", "Error de API o red.", "Reintentar y verificar el estado del servicio."],
        ]
    )
    manual.heading("14.1 Qu\u00e9 hacer ante un pago incierto", 2)
    manual.bullets(
        [
            "Verifique el estado de la transacci\u00f3n en Ventas y Tickets.",
            "Consulte el detalle del ticket para confirmar el monto.",
            "Si el cliente reporta un cobro sin ticket, revise el historial de transacciones.",
            "Contacte al \u00e1rea de TI si persisten las discrepancias.",
        ],
        ordered=True,
    )

    manual.heading("15. Seguridad, privacidad y buenas pr\u00e1cticas", 1)
    manual.bullets(
        [
            "Use una contrase\u00f1a exclusiva y no la comparta.",
            "Compruebe la URL antes de ingresar credenciales.",
            "Cierre sesi\u00f3n en equipos compartidos o p\u00fablicos.",
            "No comparta informaci\u00f3n de tickets, transacciones o datos personales de usuarios.",
            "Revise las solicitudes de reembolso antes de aprobarlas.",
            "Evite actualizar o retroceder mientras se procesa una operaci\u00f3n cr\u00edtica.",
            "Mantenga el navegador actualizado.",
            "Reporte actividades sospechosas o accesos no autorizados.",
        ]
    )

    manual.heading("16. Criterios de aceptaci\u00f3n del administrador", 1)
    manual.table(
        [
            ["ID", "Criterio verificable", "Evidencia"],
            ["CA-01", "El administrador puede iniciar sesi\u00f3n con credenciales v\u00e1lidas.", "Figura 1"],
            ["CA-02", "La barra lateral permite navegar entre todos los m\u00f3dulos.", "Figura 2"],
            ["CA-03", "El dashboard muestra estad\u00edsticas y gr\u00e1ficos filtrables por per\u00edodo.", "Figuras 3 y 4"],
            ["CA-04", "El cat\u00e1logo lista y filtra pel\u00edculas.", "Figuras 5 y 6"],
            ["CA-05", "El administrador puede gestionar cines y salas (CRUD).", "Figuras 7 a 11"],
            ["CA-06", "La programaci\u00f3n permite crear y visualizar funciones por fecha.", "Figuras 12, 13 y 14"],
            ["CA-07", "Las ventas se listan con filtros y detalle de ticket.", "Figuras 15, 16 y 17"],
            ["CA-08", "El administrador puede gestionar el flujo completo de reembolsos.", "Figuras 18 a 23"],
            ["CA-09", "La validaci\u00f3n de entrada permite escanear y verificar tickets.", "Figuras 24 y 25"],
            ["CA-10", "El m\u00f3dulo de ayuda permite gestionar consultas de usuarios.", "Figura 26"],
        ]
    )

    manual.heading("17. Gesti\u00f3n de configuraci\u00f3n del manual", 1)
    manual.paragraph(
        "Para el curso de Gesti\u00f3n de la Configuraci\u00f3n de Software, este manual se considera un elemento de "
        "configuraci\u00f3n documental relacionado con el frontend administrativo."
    )
    manual.heading("17.1 Elementos bajo control", 2)
    manual.table(
        [
            ["Elemento", "Identificador / ubicaci\u00f3n", "Disparador de actualizaci\u00f3n"],
            ["Manual editable", "docs-manual-admin/Manual_de_Administrador_FILMATE.docx", "Cambio funcional o correcci\u00f3n aprobada."],
            ["Manual portable", "docs-manual-admin/Manual_de_Administrador_FILMATE.pdf", "Nueva versi\u00f3n del manual editable."],
            ["Fuente trazable", "docs-manual-admin/Manual_de_Administrador_FILMATE.md", "Cambio de contenido o estructura."],
            ["Evidencias", "docs-manual-admin/capturas/*.png", "Cambio visual, etiqueta, ruta o resultado."],
            ["Automatizaci\u00f3n", "scripts/manual/generate_admin_manual.py", "Cambio de contenido o estructura del manual."],
        ]
    )
    manual.heading("17.2 Procedimiento de cambio", 2)
    manual.bullets(
        [
            "Registrar la solicitud de cambio y el motivo.",
            "Identificar requisitos, pantallas y procedimientos afectados.",
            "Actualizar la aplicaci\u00f3n o el contrato API correspondiente.",
            "Ejecutar nuevamente las capturas afectadas.",
            "Actualizar contenido, control de versiones y matriz de trazabilidad.",
            "Revisar ortograf\u00eda, consistencia, enlaces e im\u00e1genes.",
            "Generar Word y PDF desde la fuente.",
            "Aprobar y etiquetar la nueva l\u00ednea base documental.",
        ],
        ordered=True,
    )
    manual.heading("17.3 Nomenclatura de versiones", 2)
    manual.paragraph(
        "Se recomienda usar versi\u00f3n mayor cuando cambia el flujo o alcance del manual; versi\u00f3n menor cuando se "
        "agregan funciones compatibles; y revisi\u00f3n de parche cuando solo se corrigen redacci\u00f3n, formato o capturas sin cambiar el procedimiento."
    )

    manual.heading("18. Glosario", 1)
    manual.table(
        [
            ["T\u00e9rmino", "Definici\u00f3n"],
            ["API", "Servicio que comunica el frontend con datos y operaciones del sistema."],
            ["Dashboard", "Panel principal con m\u00e9tricas y gr\u00e1ficos del negocio."],
            ["CRUD", "Crear, Leer, Actualizar, Eliminar: operaciones b\u00e1sicas de gesti\u00f3n."],
            ["Funci\u00f3n", "Exhibici\u00f3n de una pel\u00edcula en fecha, hora, cine y sala determinados."],
            ["Sala", "Espacio f\u00edsico donde se exhiben pel\u00edculas."],
            ["Reembolso", "Devoluci\u00f3n del monto pagado por un ticket."],
            ["QR", "C\u00f3digo gr\u00e1fico asociado al ticket para validaci\u00f3n de entrada."],
            ["Ticket", "Comprobante de compra con informaci\u00f3n de la funci\u00f3n y asientos."],
            ["L\u00ednea base", "Versi\u00f3n aprobada y controlada de software, datos o documentaci\u00f3n."],
            ["Elemento de configuraci\u00f3n", "Activo sujeto a identificaci\u00f3n, versi\u00f3n, cambio y auditor\u00eda."],
        ]
    )

    manual.heading("19. Gu\u00eda r\u00e1pida", 1)
    manual.paragraph(
        "Gestionar pel\u00edculas: Cat\u00e1logo \u2192 Agregar / Editar / Eliminar pel\u00edcula.",
        callout=True,
    )
    manual.paragraph(
        "Programar funci\u00f3n: Programaci\u00f3n \u2192 Seleccionar fecha \u2192 Agregar funci\u00f3n \u2192 Elegir pel\u00edcula, sala y horario.",
        callout=True,
    )
    manual.paragraph(
        "Procesar reembolso: Ventas \u2192 Ticket \u2192 Solicitar reembolso \u2192 Revisar \u2192 Aprobar o rechazar.",
        callout=True,
    )
    manual.paragraph(
        "Validar entrada: Validaci\u00f3n de Entrada \u2192 Escanear QR \u2192 Verificar \u2192 Confirmar ingreso.",
        callout=True,
    )


def build_manual() -> None:
    manual = Manual()
    manual.cover()
    add_control_and_scope(manual)
    add_access_and_sidebar(manual)
    add_dashboard(manual)
    add_catalog(manual)
    add_cinemas_and_rooms(manual)
    add_scheduling(manual)
    add_sales_and_tickets(manual)
    add_refunds(manual)
    add_validation(manual)
    add_support(manual)
    add_support_and_gcs(manual)
    manual.save()


if __name__ == "__main__":
    build_manual()
